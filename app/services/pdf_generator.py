import requests
import os
import subprocess
import uuid
import shutil
import json
import time
import tempfile
import datetime
import glob
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from enum import IntEnum
from contextlib import contextmanager
from io import BytesIO

from PIL import Image
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.text.paragraph import Paragraph

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


# Image optimization settings
MAX_IMAGE_WIDTH = 1920
MAX_IMAGE_HEIGHT = 1920
IMAGE_JPEG_QUALITY = 85


def _optimize_image_stream(image_data: bytes) -> BytesIO:
    """
    Downsize and re-compress image data to reduce PDF file size.

    - Converts RGBA/P mode images to RGB (required for JPEG)
    - Downsizes images exceeding MAX_IMAGE_WIDTH x MAX_IMAGE_HEIGHT using LANCZOS
    - Saves as JPEG with quality=IMAGE_JPEG_QUALITY, optimize=True
    - Falls back to the original bytes on any error

    Args:
        image_data: Raw image bytes

    Returns:
        BytesIO stream of the optimized (or original) image
    """
    try:
        img = Image.open(BytesIO(image_data))

        # Convert palette or RGBA to RGB for JPEG output
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        # Downsize if larger than max dimensions
        img.thumbnail(
            (MAX_IMAGE_WIDTH, MAX_IMAGE_HEIGHT), Image.LANCZOS
        )

        output = BytesIO()
        img.save(output, format="JPEG", quality=IMAGE_JPEG_QUALITY, optimize=True)
        output.seek(0)
        return output
    except Exception as e:
        logger.warning(f"Image optimization failed, using original: {e}")
        return BytesIO(image_data)


# Constants
class Constants:
    """Application constants."""

    # List formatting
    BASE_LIST_INDENT = 0.25  # Base indentation for lists (inches)
    LIST_INDENT_STEP = 0.25  # Additional indent per nesting level (inches)
    LIST_INDENT_STEP_2 = 0.08  # Additional indent per nesting level (inches)
    HANGING_INDENT = 0.2  # Hanging indent for list items (inches)

    # Placeholder patterns
    PLACEHOLDER_PATTERN = r"\{\{(\w+)\}\}"
    PLACEHOLDER_FORMAT = "{{{{{}}}}}"

    # Data types
    PRE_FORMATTED_TYPE = "pre-formatted"
    DATA_TYPE_KEY = "data-type"

    # Styling
    DEFAULT_BULLET = "• "
    ORDERED_BULLET_FORMAT = "{}. "

    # Border settings
    BORDER_SIZE = "4"
    BORDER_SPACE = "6"
    BORDER_COLOR = "auto"
    BORDER_TYPE = "single"

    # Colors
    COLOR_GRAY = RGBColor(128, 128, 128)
    COLOR_RED = RGBColor(255, 0, 0)

    # Timeouts
    PDF_CONVERSION_TIMEOUT = 60
    IMAGE_DOWNLOAD_TIMEOUT = 10

    # Run formatting mappings
    RUN_FORMATTING_MAP = {
        "bold": "bold",
        "italic": "italic",
        "underline": "underline",
        "strikethrough": ("font.strike", True),
        "highlight_color": ("font.highlight_color", 7),  # Yellow
    }

    # Style fallback mappings
    STYLE_FALLBACK_MAP = {
        "strong": ("bold", True),
        "bold": ("bold", True),
        "emphasis": ("italic", True),
        "italic": ("italic", True),
        "heading": ("bold", True),
        "title": ("bold", True),
        "subtitle": ("italic", True),
    }

    # Paragraph alignment mappings
    ALIGNMENT_MAP = {"center": "CENTER", "right": "RIGHT", "left": "LEFT"}


class StyleType(IntEnum):
    """Word document style types."""

    PARAGRAPH = 1
    CHARACTER = 2
    TABLE = 3
    LIST = 4


class BlockType:
    """Supported block types for document rendering."""

    PARAGRAPH = "paragraph"
    HEADING = "heading"
    LIST = "list"
    TABLE = "table"
    IMAGE = "image"
    SECTION = "section"
    DIVIDER = "divider"
    SECTION_TITLE = "sectiontitle"


def generate_pdf(
    template_path: str,
    json_data: Dict[str, Any],
    output_dir: str,
    output_filename: str = None,
) -> str:
    """
    Generate PDF from template and JSON data (simple version).

    Args:
        template_path: Path to the Word template file
        json_data: Data to inject into the template
        output_dir: Directory to save output files
        output_filename: Base name for output files (without extension). If None, generates a unique filename.

    Returns:
        str: Path to generated PDF file
    """
    # Generate unique filename if not provided
    if output_filename is None:
        output_filename = f"output_{uuid.uuid4().hex[:8]}"

    temp_dir = "temp_images"
    os.makedirs(temp_dir, exist_ok=True)

    try:
        docx_path = os.path.join(output_dir, f"{output_filename}.docx")
        pdf_path = os.path.join(output_dir, f"{output_filename}.pdf")

        inject_data_into_docx(template_path, json_data, docx_path, temp_dir)
        convert_to_pdf(docx_path, pdf_path)
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

    return pdf_path


def generate_pdf_dynamic(
    template_path: str,
    json_data: Dict[str, Any],
    output_dir: str,
    output_filename: str = "output",
) -> Dict[str, Any]:
    """
    Generate PDF using DocxRenderer and DocumentProcessor classes for dynamic content.

    This function is designed for API usage and provides comprehensive cleanup of temporary files.

    Args:
        template_path: Path to the Word template file
        json_data: Data containing both static and dynamic content
        output_dir: Directory to save the output files
        output_filename: Base name for output files (without extension)

    Returns:
        Dict containing paths to generated files and status information:
        {
            'success': bool,
            'docx_path': str,
            'pdf_path': str,
            'error': str (if any),
            'processing_time': float (seconds)
        }
    """
    start_time = time.time()
    temp_files = []
    temp_dirs = []

    try:
        _validate_inputs(template_path, json_data, output_dir, output_filename)
        os.makedirs(output_dir, exist_ok=True)

        processor = DocumentProcessor(template_path, output_dir)
        docx_path, pdf_path = processor.process_document(
            json_data, output_filename
        )

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file was not created: {docx_path}")

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file was not created: {pdf_path}")

        processing_time = time.time() - start_time

        return {
            "success": True,
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "error": None,
            "processing_time": round(processing_time, 2),
        }

    except (FileNotFoundError, PermissionError, Exception) as e:
        processing_time = time.time() - start_time
        error_msg = f"{type(e).__name__}: {str(e)}"

        return {
            "success": False,
            "docx_path": None,
            "pdf_path": None,
            "error": error_msg,
            "processing_time": round(processing_time, 2),
        }

    finally:
        _cleanup_temp_files(temp_files, temp_dirs)


def get_template_styles(template_path: str) -> Dict[str, List[str]]:
    """
    Extract and return all available styles from a Word template.

    Args:
        template_path: Path to the Word template file

    Returns:
        Dict containing categorized styles:
        {
            'paragraph': List[str],
            'character': List[str],
            'table': List[str],
            'numbering': List[str]
        }
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    try:
        document = Document(template_path)

        style_groups = {
            "paragraph": [],
            "character": [],
            "table": [],
            "numbering": [],
        }

        # WD_STYLE_TYPE constants
        STYLE_TYPE_MAPPING = {
            1: "paragraph",  # WD_STYLE_TYPE.PARAGRAPH
            2: "character",  # WD_STYLE_TYPE.CHARACTER
            3: "table",  # WD_STYLE_TYPE.TABLE
            4: "numbering",  # WD_STYLE_TYPE.LIST
        }

        for style in document.styles:
            style_category = STYLE_TYPE_MAPPING.get(style.type)
            if style_category:
                style_groups[style_category].append(style.name)

        # Sort styles alphabetically
        for category in style_groups:
            style_groups[category].sort()

        return style_groups

    except Exception as e:
        raise Exception(f"Error reading template styles: {e}")


def print_template_styles(template_path: str) -> None:
    """
    Print all available styles in the Word template in a formatted way.

    Args:
        template_path: Path to the Word template file
    """
    try:
        style_groups = get_template_styles(template_path)

        style_info = [
            (
                "PARAGRAPH STYLES",
                "paragraph",
                "Use these for block-level 'style' property",
            ),
            (
                "CHARACTER STYLES",
                "character",
                "Use these for run-level 'style' property within 'runs'",
            ),
            ("TABLE STYLES", "table", "Use these for table styling"),
            ("LIST STYLES", "numbering", "Use these for list styling"),
        ]

        lines = ["AVAILABLE STYLES IN TEMPLATE"]
        for title, key, description in style_info:
            styles = style_groups[key]
            if styles:
                lines.append(f"{title} ({len(styles)}): ({description})")
                for style in styles:
                    lines.append(f"   • {style}")

        total_styles = sum(len(styles) for styles in style_groups.values())
        lines.append(f"TOTAL STYLES: {total_styles}")
        logger.debug("\n".join(lines))

    except Exception as e:
        logger.debug("Error reading template styles: %s", e)


def get_pdf_info(pdf_path: str) -> Dict[str, Any]:
    """
    Get information about a generated PDF file.

    Args:
        pdf_path: Path to the PDF file

    Returns:
        Dict containing PDF file information
    """
    if not os.path.exists(pdf_path):
        return {
            "exists": False,
            "size_bytes": 0,
            "size_mb": 0.0,
            "created_time": None,
            "modified_time": None,
        }

    stat = os.stat(pdf_path)
    size_bytes = stat.st_size
    size_mb = round(size_bytes / (1024 * 1024), 2)

    created_time = datetime.datetime.fromtimestamp(stat.st_ctime).isoformat()
    modified_time = datetime.datetime.fromtimestamp(stat.st_mtime).isoformat()

    return {
        "exists": True,
        "size_bytes": size_bytes,
        "size_mb": size_mb,
        "created_time": created_time,
        "modified_time": modified_time,
    }


def cleanup_old_files(
    output_dir: str, max_age_hours: int = 24
) -> Dict[str, Any]:
    """
    Clean up old generated files from the output directory.

    Args:
        output_dir: Directory to clean up
        max_age_hours: Maximum age of files in hours before cleanup

    Returns:
        Dict containing cleanup results
    """
    if not os.path.exists(output_dir):
        return {"files_removed": 0, "space_freed_mb": 0.0, "errors": []}

    current_time = time.time()
    max_age_seconds = max_age_hours * 3600
    files_removed = 0
    space_freed = 0
    errors = []

    patterns = [
        os.path.join(output_dir, "*.pdf"),
        os.path.join(output_dir, "*.docx"),
    ]

    for pattern in patterns:
        for file_path in glob.glob(pattern):
            try:
                if os.path.getmtime(file_path) < current_time - max_age_seconds:
                    file_size = os.path.getsize(file_path)
                    os.unlink(file_path)
                    files_removed += 1
                    space_freed += file_size
            except OSError as e:
                errors.append(f"Error removing {file_path}: {e}")

    space_freed_mb = round(space_freed / (1024 * 1024), 2)

    return {
        "files_removed": files_removed,
        "space_freed_mb": space_freed_mb,
        "errors": errors,
    }


def inject_data_into_docx(
    template_path: str,
    json_data: Dict[str, Any],
    output_docx: str,
    temp_dir: str,
) -> None:
    """
    Inject JSON data into a DOCX template, including image processing.

    Args:
        template_path: Path to the Word template file
        json_data: Data to inject into the template
        output_docx: Path for the output DOCX file
        temp_dir: Directory for temporary files
    """
    doc = DocxTemplate(template_path)
    temp_images = []

    def process_images(obj: Any) -> None:
        """Recursively process images in the data structure."""
        if isinstance(obj, dict):
            for key in list(obj.keys()):
                if isinstance(obj[key], (dict, list)):
                    process_images(obj[key])
                elif key.endswith("image_url") and obj[key]:
                    _process_single_image(obj, key, temp_dir, temp_images, doc)
        elif isinstance(obj, list):
            for item in obj:
                process_images(item)

    process_images(json_data)
    doc.render(json_data)

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    doc.save(output_docx)


def convert_to_pdf(input_docx: str, output_pdf: str) -> None:
    """
    Convert a DOCX file to PDF using UnoServer.

    Args:
        input_docx: Path to the input DOCX file
        output_pdf: Path for the output PDF file
    """
    from app.services.unoserver_converter import UnoServerConverter

    converter = UnoServerConverter()

    # Check if UnoServer is available
    if not converter.is_available():
        raise Exception(
            "UnoServer is not available. "
            "Ensure the unoserver container is running."
        )

    # Convert using UnoServer
    output_dir = os.path.dirname(output_pdf)
    converter.convert_to_pdf(input_docx, output_dir)


def load_data(file_path: str) -> Dict[str, Any]:
    """
    Load JSON data from a file.

    Args:
        file_path: Path to the JSON file

    Returns:
        Dict containing the loaded data

    Raises:
        Exception: If file is not found or contains invalid JSON
    """
    file_path = Path(file_path)

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        raise Exception(f"Data file not found: {file_path}")
    except json.JSONDecodeError as e:
        raise Exception(f"Invalid JSON in data file: {e}")


def prepare_image_data(
    template: DocxTemplate, data: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Process image URLs in data and convert them to InlineImage objects.

    Looks for keys ending with '_image_url' and converts valid URLs to InlineImage.

    Args:
        template: DocxTemplate instance
        data: Dictionary containing the data

    Returns:
        Modified data dictionary with InlineImage objects
    """
    processed_data = data.copy()

    # Find and process all image fields ending with _image_url
    image_keys = [key for key in data.keys() if key.endswith("_image_url")]

    for key in image_keys:
        image_url = data[key]

        # Skip if empty or not a string
        if not image_url or not isinstance(image_url, str):
            continue

        # Only process valid URLs
        if not image_url.startswith(("http://", "https://")):
            continue

        try:
            # Download and create InlineImage
            response = requests.get(
                image_url, timeout=Constants.IMAGE_DOWNLOAD_TIMEOUT
            )
            response.raise_for_status()
            image_stream = _optimize_image_stream(response.content)
            processed_data[key] = InlineImage(
                template, image_stream, width=Inches(2)
            )
            logger.info(f"Processed image: {key}")
        except Exception as e:
            logger.warning(f"Could not process image '{key}': {e}")
            # Keep original value, don't set to None

    # Recursively process nested structures
    for key, value in data.items():
        if isinstance(value, dict) and Constants.DATA_TYPE_KEY not in value:
            processed_data[key] = prepare_image_data(template, value)
        elif isinstance(value, list):
            processed_data[key] = [
                prepare_image_data(template, item)
                if isinstance(item, dict)
                else item
                for item in value
            ]

    return processed_data


# Private helper functions


def _validate_inputs(
    template_path: str,
    json_data: Dict[str, Any],
    output_dir: str,
    output_filename: str,
) -> None:
    """Validate input parameters for generate_pdf_dynamic function."""
    if not template_path or not isinstance(template_path, str):
        raise ValueError("template_path must be a non-empty string")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    if not template_path.lower().endswith((".docx", ".doc")):
        raise ValueError(
            "Template file must be a Word document (.docx or .doc)"
        )

    if not json_data or not isinstance(json_data, dict):
        raise ValueError("json_data must be a non-empty dictionary")

    if not output_dir or not isinstance(output_dir, str):
        raise ValueError("output_dir must be a non-empty string")

    if not output_filename or not isinstance(output_filename, str):
        raise ValueError("output_filename must be a non-empty string")

    invalid_chars = '<>:"/\\|?*'
    if any(char in output_filename for char in invalid_chars):
        raise ValueError(
            f"output_filename contains invalid characters: {invalid_chars}"
        )


def _cleanup_temp_files(temp_files: List[str], temp_dirs: List[str]) -> None:
    """Clean up temporary files and directories."""
    for temp_file in temp_files:
        try:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        except OSError:
            pass  # Ignore cleanup errors

    for temp_dir in temp_dirs:
        try:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except OSError:
            pass  # Ignore cleanup errors

    _cleanup_orphaned_temp_files()


def _cleanup_orphaned_temp_files() -> None:
    """Clean up orphaned temporary files in common temp locations."""
    current_time = time.time()

    temp_patterns = [
        os.path.join(tempfile.gettempdir(), "tmp*.docx"),
        os.path.join(tempfile.gettempdir(), "tmp*.jpg"),
        os.path.join(tempfile.gettempdir(), "tmp*.png"),
        os.path.join(tempfile.gettempdir(), "tmp*.gif"),
    ]

    for pattern in temp_patterns:
        try:
            for temp_file in glob.glob(pattern):
                if (
                    os.path.getmtime(temp_file) < current_time - 3600
                ):  # 1 hour old
                    try:
                        os.unlink(temp_file)
                    except OSError:
                        pass
        except Exception:
            pass

    local_temp_dirs = ["temp_images", "temp_files", "temp_docs"]
    for temp_dir in local_temp_dirs:
        try:
            if os.path.exists(temp_dir) and os.path.isdir(temp_dir):
                if (
                    not os.listdir(temp_dir)
                    or os.path.getmtime(temp_dir) < current_time - 3600
                ):
                    shutil.rmtree(temp_dir)
        except OSError:
            pass


def _process_single_image(
    obj: Dict[str, Any],
    key: str,
    temp_dir: str,
    temp_images: List[str],
    doc: DocxTemplate,
) -> None:
    """Process a single image URL and convert to InlineImage."""
    image_path = None
    try:
        response = requests.get(obj[key], stream=True)
        response.raise_for_status()

        image_path = os.path.join(temp_dir, f"{uuid.uuid4()}.jpg")
        temp_images.append(image_path)

        optimized = _optimize_image_stream(response.content)
        with open(image_path, "wb") as img_file:
            img_file.write(optimized.read())

        inline_image = InlineImage(doc, image_path)
        obj[key] = inline_image

    except Exception as e:
        obj[f"{key}_error"] = f"Error processing image: {str(e)}"
        if image_path and os.path.exists(image_path):
            temp_images.append(image_path)


class PlaceholderReplacer:
    """
    Handles replacement of placeholders in Word documents while preserving formatting.

    This class manages the complex task of finding and replacing placeholder text
    throughout a Word document, including in headers, footers, and tables.
    """

    def __init__(self, document: Document):
        """
        Initialize the PlaceholderReplacer.

        Args:
            document: The Word document to process
        """
        self.document = document

    def replace_all_placeholders(self, data: Dict[str, Any]) -> None:
        """
        Replace all placeholders in the document with data values.

        Args:
            data: Dictionary mapping placeholder names to replacement values
        """
        # Separate static data from pre-formatted content
        static_data = self._extract_static_data(data)

        # Process all document sections
        self._process_document_content(static_data)
        self._process_headers(static_data)
        self._process_footers(static_data)

    def _extract_static_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract non-preformatted data for placeholder replacement."""
        static_data = {}
        for key, value in data.items():
            if not self._is_preformatted(value):
                static_data[key] = value
        return static_data

    @staticmethod
    def _is_preformatted(value: Any) -> bool:
        """Check if a value is pre-formatted content."""
        return (
            isinstance(value, dict)
            and value.get(Constants.DATA_TYPE_KEY)
            == Constants.PRE_FORMATTED_TYPE
        )

    def _process_document_content(self, static_data: Dict[str, Any]) -> None:
        """Process main document paragraphs and tables."""
        # Process paragraphs
        for paragraph in self.document.paragraphs:
            self._replace_paragraph_placeholders(paragraph, static_data)

        # Process tables
        for table in self.document.tables:
            self._process_table(table, static_data)

    def _process_table(self, table, static_data: Dict[str, Any]) -> None:
        """Process all cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_paragraph_placeholders(paragraph, static_data)

    def _process_headers(self, static_data: Dict[str, Any]) -> None:
        """Process all document headers."""
        for section in self.document.sections:
            if section.header:
                self._process_header_footer(section.header, static_data)

    def _process_footers(self, static_data: Dict[str, Any]) -> None:
        """Process all document footers."""
        for section in self.document.sections:
            if section.footer:
                self._process_header_footer(section.footer, static_data)

    def _process_header_footer(
        self, header_footer, static_data: Dict[str, Any]
    ) -> None:
        """Process paragraphs and tables in headers/footers."""
        # Process paragraphs
        for paragraph in header_footer.paragraphs:
            self._replace_paragraph_placeholders(paragraph, static_data)

        # Process tables
        for table in header_footer.tables:
            self._process_table(table, static_data)

    def _replace_paragraph_placeholders(
        self, paragraph: Paragraph, data: Dict[str, Any]
    ) -> None:
        """Replace placeholders in a paragraph while preserving formatting."""
        full_text = self._get_paragraph_text(paragraph)

        # Find all placeholders
        placeholders = re.findall(Constants.PLACEHOLDER_PATTERN, full_text)

        if not placeholders:
            return

        # Process each placeholder
        for placeholder in placeholders:
            if placeholder in data:
                placeholder_text = Constants.PLACEHOLDER_FORMAT.format(
                    placeholder
                )
                replacement_text = str(data[placeholder])
                self._replace_placeholder_in_paragraph(
                    paragraph, placeholder_text, replacement_text
                )

    @staticmethod
    def _get_paragraph_text(paragraph: Paragraph) -> str:
        """Get the complete text of a paragraph from all runs."""
        return "".join(run.text for run in paragraph.runs)

    def _replace_placeholder_in_paragraph(
        self, paragraph: Paragraph, placeholder: str, replacement: str
    ) -> None:
        """Replace a specific placeholder in a paragraph."""
        runs = paragraph.runs

        # Preserve paragraph formatting
        original_format = self._preserve_paragraph_format(paragraph)

        # Build character-to-run mapping
        char_to_run = self._build_char_to_run_map(runs)

        # Find placeholder position
        full_text = self._get_paragraph_text(paragraph)
        placeholder_start = full_text.find(placeholder)

        if placeholder_start == -1 or not char_to_run:
            return

        placeholder_end = placeholder_start + len(placeholder)

        # Determine affected runs
        start_run_idx = char_to_run[
            min(placeholder_start, len(char_to_run) - 1)
        ]
        end_run_idx = char_to_run[
            min(placeholder_end - 1, len(char_to_run) - 1)
        ]

        # Calculate positions within runs
        chars_before_start = sum(
            len(runs[i].text) for i in range(start_run_idx)
        )
        start_pos_in_run = placeholder_start - chars_before_start

        chars_before_end = sum(len(runs[i].text) for i in range(end_run_idx))
        end_pos_in_run = placeholder_end - chars_before_end

        # Perform replacement
        self._perform_replacement(
            runs,
            start_run_idx,
            end_run_idx,
            start_pos_in_run,
            end_pos_in_run,
            replacement,
        )

        # Restore formatting
        self._restore_paragraph_format(paragraph, original_format)

    @staticmethod
    def _build_char_to_run_map(runs: List) -> List[int]:
        """Build a map from character positions to run indices."""
        char_to_run = []
        for i, run in enumerate(runs):
            char_to_run.extend([i] * len(run.text))
        return char_to_run

    @staticmethod
    def _perform_replacement(
        runs: List,
        start_run_idx: int,
        end_run_idx: int,
        start_pos: int,
        end_pos: int,
        replacement: str,
    ) -> None:
        """Perform the actual text replacement in runs."""
        if start_run_idx == end_run_idx:
            # Placeholder within single run
            run = runs[start_run_idx]
            original_text = run.text
            run.text = (
                original_text[:start_pos]
                + replacement
                + original_text[end_pos:]
            )
        else:
            # Placeholder spans multiple runs
            runs[start_run_idx].text = (
                runs[start_run_idx].text[:start_pos] + replacement
            )

            # Clear middle runs
            for i in range(start_run_idx + 1, end_run_idx):
                runs[i].text = ""

            # Adjust last run
            if end_run_idx < len(runs):
                runs[end_run_idx].text = runs[end_run_idx].text[end_pos:]

    def _preserve_paragraph_format(
        self, paragraph: Paragraph
    ) -> Dict[str, Any]:
        """Preserve paragraph formatting properties."""
        para_format = paragraph.paragraph_format
        return {
            "alignment": para_format.alignment,
            "left_indent": para_format.left_indent,
            "right_indent": para_format.right_indent,
            "first_line_indent": para_format.first_line_indent,
            "line_spacing": para_format.line_spacing,
            "line_spacing_rule": para_format.line_spacing_rule,
            "space_before": para_format.space_before,
            "space_after": para_format.space_after,
            "keep_together": para_format.keep_together,
            "keep_with_next": para_format.keep_with_next,
            "page_break_before": para_format.page_break_before,
            "widow_control": para_format.widow_control,
        }

    def _restore_paragraph_format(
        self, paragraph: Paragraph, format_dict: Dict[str, Any]
    ) -> None:
        """Restore paragraph formatting properties."""
        para_format = paragraph.paragraph_format

        try:
            # Only set non-None values to preserve template defaults
            for key, value in format_dict.items():
                if value is not None:
                    setattr(para_format, key, value)
        except Exception as e:
            # Continue even if formatting fails
            logger.warning(f"Could not restore paragraph formatting: {e}")


class DocxRenderer:
    """
    Renders Word document content from a declarative, block-based JSON structure.

    This class handles the conversion of JSON-defined document structures into
    properly formatted Word document elements, including paragraphs, lists,
    tables, images, and more.
    """

    def __init__(
        self,
        document: Document,
        insert_at_parent: Optional[Any] = None,
        insert_at_index: Optional[int] = None,
    ):
        """
        Initialize the DocxRenderer.

        Args:
            document: The Word document object to render into
            insert_at_parent: Optional parent element for positioned insertion
            insert_at_index: Optional index for positioned insertion
        """
        self.document = document
        self.current_list_level = 0
        self.insert_at_parent = insert_at_parent
        self.insert_at_index = insert_at_index
        self._current_insert_index = (
            insert_at_index if insert_at_index is not None else 0
        )

    def render_document(self, data: Union[List, Dict]) -> None:
        """
        Render the entire document from JSON data.

        Args:
            data: Either a list of blocks or a dictionary with document structure
        """
        # Handle both legacy (with document wrapper) and new format (direct array)
        if isinstance(data, list):
            body_blocks = data
        else:
            doc_data = data.get("document", {})
            body_blocks = doc_data.get("body", [])

        for block in body_blocks:
            self.render_block(block)

    def _create_paragraph_at_position(self) -> Paragraph:
        """
        Create a paragraph at the specified insertion position.

        Returns:
            Newly created Paragraph object
        """
        if self.insert_at_parent is not None:
            from docx.oxml.parser import parse_xml

            p_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            p_element = parse_xml(p_xml)
            self.insert_at_parent.insert(self._current_insert_index, p_element)
            self._current_insert_index += 1
            return Paragraph(p_element, self.document)
        else:
            return self.document.add_paragraph()

    def _create_heading_at_position(self, level: int = 1) -> Paragraph:
        """
        Create a heading at the specified insertion position.

        Args:
            level: Heading level (1-9)

        Returns:
            Newly created heading paragraph
        """
        if self.insert_at_parent is not None:
            p = self._create_paragraph_at_position()
            try:
                p.style = f"Heading {level}"
            except KeyError:
                p.style = "Heading 1"  # fallback to Heading 1
                logger.warning(
                    f"Heading {level} style not found, using Heading 1"
                )
            return p
        else:
            return self.document.add_heading(level=level)

    def _create_table_at_position(
        self, rows: int, cols: int, style: Optional[str] = None
    ):
        """
        Create a table at the specified insertion position.

        Args:
            rows: Number of rows
            cols: Number of columns
            style: Optional table style name

        Returns:
            Newly created Table object
        """
        if self.insert_at_parent is not None:
            from docx.table import Table
            from docx.oxml.parser import parse_xml

            # Create table XML structure
            tbl_xml = self._generate_table_xml(cols)
            tbl_element = parse_xml(tbl_xml)

            # Add rows
            for _ in range(rows):
                tr_xml = self._generate_table_row_xml(cols)
                tr_element = parse_xml(tr_xml)
                tbl_element.append(tr_element)

            self.insert_at_parent.insert(
                self._current_insert_index, tbl_element
            )
            self._current_insert_index += 1
            table = Table(tbl_element, self.document)

            # Apply style if provided
            self._apply_table_style(table, style)
            return table
        else:
            # Standard table creation
            table = self.document.add_table(rows=rows, cols=cols)
            self._apply_table_style(table, style)
            return table

    @staticmethod
    def _generate_table_xml(cols: int) -> str:
        """Generate XML for table structure."""
        return f"""
        <w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tblPr>
                <w:tblW w:w="0" w:type="auto"/>
            </w:tblPr>
            <w:tblGrid>
                {"".join("<w:gridCol/>" for _ in range(cols))}
            </w:tblGrid>
        </w:tbl>
        """

    @staticmethod
    def _generate_table_row_xml(cols: int) -> str:
        """Generate XML for table row."""
        return f"""
        <w:tr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            {"".join("<w:tc><w:tcPr></w:tcPr><w:p></w:p></w:tc>" for _ in range(cols))}
        </w:tr>
        """

    def _apply_table_style(self, table, style: Optional[str]) -> None:
        """Apply style to a table with error handling."""
        if style:
            try:
                table.style = style
            except KeyError:
                logger.warning(
                    f"Table style '{style}' not found. Using default."
                )

    def render_block(self, block: Dict[str, Any]) -> None:
        """Render a single block based on its type."""
        block_type = block.get("type", "Paragraph").lower()
        style = block.get("style")
        data = block.get("data", {})

        render_method = getattr(
            self, f"_render_{block_type}", self._render_unsupported
        )
        render_method(data, style)

    def _render_heading(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Render a heading element."""
        text = data.get("text", "")

        # Extract level from style name if present
        level = self._extract_heading_level(style)

        if self.insert_at_parent is not None:
            p = self._create_heading_at_position(level=level)
        else:
            p = self.document.add_heading(level=level)

        p.add_run(text)

        # Apply custom style if provided
        if style:
            self._apply_paragraph_style(p, style)

    @staticmethod
    def _extract_heading_level(style: Optional[str]) -> int:
        """Extract heading level from style name."""
        if style and style.startswith("Heading"):
            try:
                return int(style[-1])
            except (ValueError, IndexError):
                pass
        return 1

    def _render_sectiontitle(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Render a section title with proper spacing and underline."""
        text = data.get("text", "")

        p = self._create_paragraph_at_position()
        run = p.add_run(text)

        try:
            p.style = "Heading 2"
        except KeyError:
            run.bold = True
            logger.warning(
                "'Heading 2' style not found, using fallback formatting"
            )

        self._add_bottom_border_to_paragraph(p)

    def _add_bottom_border_to_paragraph(self, paragraph: Paragraph) -> None:
        """Add a black bottom border to the given paragraph."""
        pPr = paragraph._element.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")
        bottom_bdr = self._create_border_element()
        pBdr.append(bottom_bdr)
        pPr.append(pBdr)

    @staticmethod
    def _create_border_element() -> OxmlElement:
        """Create a border XML element."""
        bottom_bdr = OxmlElement("w:bottom")
        bottom_bdr.set(qn("w:val"), Constants.BORDER_TYPE)
        bottom_bdr.set(qn("w:sz"), Constants.BORDER_SIZE)
        bottom_bdr.set(qn("w:space"), Constants.BORDER_SPACE)
        bottom_bdr.set(qn("w:color"), Constants.BORDER_COLOR)
        return bottom_bdr

    def _render_paragraph(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Render a paragraph with rich text runs."""
        p = self._create_paragraph_at_position()

        if style:
            self._apply_paragraph_style(p, style)

        is_inline = data.get("inline", False)

        # Process text runs
        for run_data in data.get("runs", []):
            text = run_data.get("text", "")
            run = p.add_run(text)

            self._apply_run_formatting(run, run_data)

            if not run_data.get("inline", True):
                run.add_break()

        # Adjust spacing for inline paragraphs
        if is_inline:
            p.paragraph_format.space_after = Pt(0)

    def _apply_run_formatting(self, run, run_data: Dict[str, Any]) -> None:
        """Apply formatting to a text run based on run_data properties."""
        # Handle hyperlinks first
        if run_data.get("link"):
            self._add_hyperlink(run, run_data["link"])

        # Apply basic formatting using Constants
        for key, value in Constants.RUN_FORMATTING_MAP.items():
            if run_data.get(key):
                if isinstance(value, tuple):
                    attr_path, attr_value = value
                    self._set_nested_attribute(run, attr_path, attr_value)
                else:
                    setattr(run, value, True)

        # Handle color
        self._apply_run_color(run, run_data.get("color"))

        # Apply style
        run_style = run_data.get("style")
        if run_style:
            self._apply_run_style(run, run_style)

    @staticmethod
    def _set_nested_attribute(obj: Any, attr_path: str, value: Any) -> None:
        """Set a nested attribute using dot notation."""
        attrs = attr_path.split(".")
        for attr in attrs[:-1]:
            obj = getattr(obj, attr)
        setattr(obj, attrs[-1], value)

    def _apply_run_color(self, run, color_str: Optional[str]) -> None:
        """Apply color to a text run."""
        if color_str and len(color_str) == 6:
            try:
                r = int(color_str[0:2], 16)
                g = int(color_str[2:4], 16)
                b = int(color_str[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                logger.warning(f"Invalid color format '{color_str}'")

    def _apply_run_style(self, run, style_name: str) -> None:
        """Apply style to a run with fallback formatting."""
        try:
            run.style = style_name
        except (KeyError, ValueError):
            logger.warning(f"Cannot apply style '{style_name}' to text run.")
            self._apply_fallback_run_formatting(run, style_name)

    def _apply_fallback_run_formatting(self, run, style_name: str) -> None:
        """Apply fallback formatting when a style cannot be applied."""
        style_lower = style_name.lower()

        for key, (attr, value) in Constants.STYLE_FALLBACK_MAP.items():
            if key in style_lower:
                setattr(run, attr, value)
                logger.info(f"Applied {key} formatting for '{style_name}'")
                return

        logger.info(f"No fallback formatting available for '{style_name}'")

    def _add_hyperlink(self, run, url: str) -> None:
        """Add a hyperlink to a text run."""
        try:
            paragraph = run._element.getparent()

            # Create hyperlink element
            hyperlink = OxmlElement("w:hyperlink")

            # Create relationship for the hyperlink
            rel = self.document.part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink.set(qn("r:id"), rel)

            # Move run content to hyperlink
            run_element = run._element
            run_element.getparent().remove(run_element)
            hyperlink.append(run_element)

            # Insert hyperlink into paragraph
            paragraph.append(hyperlink)

        except Exception as e:
            logger.warning(f"Could not create hyperlink for '{url}': {e}")

    def _render_list(self, data: Dict[str, Any], style: Optional[str]) -> None:
        """
        Render a list (bulleted or ordered) with proper indentation and nested content.

        Supports both simple format (legacy) and advanced format (new) for list items.
        """
        list_type = data.get("list_type", "bulleted")
        list_counter = 0

        # Increment nesting level
        self.current_list_level += 1

        try:
            for item in data.get("items", []):
                if "blocks" in item:
                    # New format: item with nested blocks
                    self._render_list_item_with_blocks(
                        item, list_type, list_counter, style
                    )
                else:
                    # Legacy format: item with just runs
                    self._render_simple_list_item(
                        item, list_type, list_counter, style
                    )

                if list_type == "ordered":
                    list_counter += 1
        finally:
            # Always decrement nesting level
            self.current_list_level -= 1

    def _calculate_list_indentation(self) -> Tuple[Inches, Inches]:
        """Calculate indentation based on current nesting level."""
        left_indent = (
            Constants.BASE_LIST_INDENT
            + (self.current_list_level - 1) * Constants.LIST_INDENT_STEP
        )
        return Inches(left_indent), Inches(-Constants.HANGING_INDENT)

    @staticmethod
    def _get_bullet_text(list_type: str, list_counter: int) -> str:
        """Get appropriate bullet or number text."""
        if list_type == "ordered":
            return Constants.ORDERED_BULLET_FORMAT.format(list_counter + 1)
        return Constants.DEFAULT_BULLET

    def _render_simple_list_item(
        self,
        item: Dict[str, Any],
        list_type: str,
        list_counter: int,
        style: Optional[str],
    ) -> None:
        """Render a simple list item with just text runs."""
        p = self._create_paragraph_at_position()

        if style:
            self._apply_paragraph_style(p, style)

        # Apply indentation
        left_indent, first_line_indent = self._calculate_list_indentation()
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = first_line_indent

        # Add bullet or number
        bullet_text = self._get_bullet_text(list_type, list_counter)
        p.add_run(bullet_text)

        # Add content
        for run_data in item.get("runs", []):
            text = run_data.get("text", "")
            run = p.add_run(text)
            self._apply_run_formatting(run, run_data)

            if not run_data.get("inline", True):
                run.add_break()

    def _render_list_item_with_blocks(
        self,
        item: Dict[str, Any],
        list_type: str,
        list_counter: int,
        style: Optional[str],
    ) -> None:
        """Render a list item that can contain nested blocks."""
        blocks = item.get("blocks", [])
        if not blocks:
            return

        # Render first block with bullet/number
        first_block = blocks[0]
        first_block_type = first_block.get("type", "Paragraph").lower()

        if first_block_type == BlockType.PARAGRAPH:
            self._render_first_paragraph_with_bullet(
                first_block, list_type, list_counter, style
            )
        else:
            self._render_bullet_only(list_type, list_counter, style)
            self.render_block(first_block)

        # Render remaining blocks
        if len(blocks) > 1:
            self._render_nested_blocks(blocks[1:])

    def _render_first_paragraph_with_bullet(
        self,
        first_block: Dict[str, Any],
        list_type: str,
        list_counter: int,
        style: Optional[str],
    ) -> None:
        """Render the first paragraph block with bullet/number prefix."""
        p = self._create_paragraph_at_position()

        if style:
            self._apply_paragraph_style(p, style)

        # Apply indentation
        left_indent, first_line_indent = self._calculate_list_indentation()
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = first_line_indent

        # Add bullet or number
        bullet_text = self._get_bullet_text(list_type, list_counter)
        p.add_run(bullet_text)

        # Add paragraph content
        first_block_data = first_block.get("data", {})
        is_inline = first_block_data.get("inline", False)

        for run_data in first_block_data.get("runs", []):
            text = run_data.get("text", "")
            run = p.add_run(text)
            self._apply_run_formatting(run, run_data)

            if not run_data.get("inline", True):
                run.add_break()

        if is_inline:
            p.paragraph_format.space_after = Pt(0)

    def _render_bullet_only(
        self, list_type: str, list_counter: int, style: Optional[str]
    ) -> None:
        """Render just the bullet/number for non-paragraph first blocks."""
        p = self._create_paragraph_at_position()

        if style:
            self._apply_paragraph_style(p, style)

        left_indent, first_line_indent = self._calculate_list_indentation()
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = first_line_indent

        bullet_text = self._get_bullet_text(list_type, list_counter)
        p.add_run(bullet_text)

    def _render_nested_blocks(self, blocks: List[Dict[str, Any]]) -> None:
        """Render nested blocks with proper indentation."""
        content_tracker = self._get_content_tracker()

        for block in blocks:
            self.render_block(block)

        self._apply_nested_indentation(content_tracker)

    def _get_content_tracker(self) -> Dict[str, int]:
        """Get current document content counts for tracking new additions."""
        return {
            "para_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
        }

    def _apply_nested_indentation(
        self, content_tracker: Dict[str, int]
    ) -> None:
        """Apply increased indentation to content added after tracking point."""
        nested_indent = Inches(
            Constants.BASE_LIST_INDENT
            + self.current_list_level * Constants.LIST_INDENT_STEP_2
        )

        current_para_count = len(self.document.paragraphs)
        for i in range(content_tracker["para_count"], current_para_count):
            paragraph = self.document.paragraphs[i]
            self._apply_nested_paragraph_indentation(paragraph, nested_indent)

        current_table_count = len(self.document.tables)
        for i in range(content_tracker["table_count"], current_table_count):
            table = self.document.tables[i]
            self._apply_table_indentation(table, nested_indent)

    def _apply_nested_paragraph_indentation(
        self, paragraph, nested_indent: Inches
    ) -> None:
        """Apply proper indentation to a paragraph within nested content."""
        para_text = paragraph.text.strip()

        is_list_item = para_text.startswith("•") or (
            para_text and para_text[0].isdigit() and ". " in para_text
        )

        if is_list_item:
            # Leave list items as they are; their list rendering sets correct indentation
            return
        else:
            # For regular child paragraphs, indent slightly more than the list text
            paragraph.paragraph_format.left_indent = nested_indent

    def _apply_table_indentation(self, table, nested_indent: Inches) -> None:
        """Apply indentation to all cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.left_indent = nested_indent

    def _render_table(self, data: Dict[str, Any], style: Optional[str]) -> None:
        """Render a table."""
        rows_data = data.get("rows", [])
        if not rows_data:
            return

        num_cols = len(rows_data[0]) if rows_data else 0
        table = self._create_table_at_position(
            rows=len(rows_data), cols=num_cols, style=style
        )

        # Apply style if not already set
        if style and hasattr(table, "style"):
            self._apply_table_style(table, style)

        # Fill table cells
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    self._fill_table_cell(table.cell(i, j), cell_data)

    def _fill_table_cell(self, cell, cell_data: Dict[str, Any]) -> None:
        """Fill a table cell with content."""
        cell.text = ""
        p = cell.paragraphs[0]

        for run_data in cell_data.get("runs", []):
            run = p.add_run(run_data.get("text", ""))
            self._apply_run_formatting(run, run_data)

            if not run_data.get("inline", True):
                run.add_break()

    @contextmanager
    def _temp_image_file(self, src: str):
        """Context manager for handling temporary image files."""
        image_path = None
        temp_file_created = False

        try:
            image_path = self._download_image(src)
            temp_file_created = src.startswith(("http://", "https://"))
            yield image_path
        finally:
            # Clean up temporary file
            if temp_file_created and image_path and Path(image_path).exists():
                try:
                    Path(image_path).unlink()
                except OSError:
                    pass

    def _download_image(self, src: str) -> Optional[str]:
        """Download image from URL or return local file path."""
        # Check if local file
        if not src.startswith(("http://", "https://")):
            if Path(src).exists():
                return src
            else:
                logger.warning(f"Local image file not found: {src}")
                return None

        # Download from URL
        try:
            response = requests.get(
                src, timeout=Constants.IMAGE_DOWNLOAD_TIMEOUT
            )
            response.raise_for_status()

            # Optimize and create temporary file
            optimized = _optimize_image_stream(response.content)
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".jpg"
            ) as temp_file:
                temp_file.write(optimized.read())
                return temp_file.name

        except Exception as e:
            logger.warning(f"Could not download image from {src}: {e}")
            return None

    def _render_image(self, data: Dict[str, Any], style: Optional[str]) -> None:
        """Render an image with specified dimensions."""
        src = data.get("src", "")
        width = data.get("width")  # In inches
        height = data.get("height")  # In inches
        alt_text = data.get("alt", "")
        alignment = data.get("alignment", "left")

        if not src:
            logger.warning("Image src is required")
            return

        with self._temp_image_file(src) as image_path:
            if not image_path:
                # Add placeholder text
                self._add_image_placeholder(
                    alt_text or src, Constants.COLOR_GRAY
                )
                return

            try:
                # Add paragraph for image
                p = self._create_paragraph_at_position()
                run = p.add_run()

                # Add image with dimensions
                self._add_image_to_run(run, image_path, width, height)

                # Set alignment
                self._set_paragraph_alignment(p, alignment)

                # Apply style if provided
                if style:
                    self._apply_paragraph_style(p, style)

            except Exception as e:
                logger.warning(f"Could not insert image from '{src}': {e}")
                self._add_image_placeholder(
                    f"Image could not be loaded: {alt_text or src}",
                    Constants.COLOR_RED,
                )

    def _add_image_to_run(
        self,
        run,
        image_path: str,
        width: Optional[float],
        height: Optional[float],
    ) -> None:
        """Add an image to a run with specified dimensions."""
        if width and height:
            run.add_picture(
                image_path, width=Inches(width), height=Inches(height)
            )
        elif width:
            run.add_picture(image_path, width=Inches(width))
        elif height:
            run.add_picture(image_path, height=Inches(height))
        else:
            run.add_picture(image_path)

    def _set_paragraph_alignment(
        self, paragraph: Paragraph, alignment: str
    ) -> None:
        """Set paragraph alignment."""
        if alignment in Constants.ALIGNMENT_MAP:
            alignment_attr = Constants.ALIGNMENT_MAP[alignment]
            paragraph.alignment = getattr(
                WD_PARAGRAPH_ALIGNMENT, alignment_attr
            )

    def _add_image_placeholder(self, text: str, color: RGBColor) -> None:
        """Add placeholder text for missing images."""
        p = self._create_paragraph_at_position()
        run = p.add_run(f"[{text}]")
        run.italic = True
        run.font.color.rgb = color

    def _apply_paragraph_style(self, paragraph: Paragraph, style: str) -> None:
        """Apply style to paragraph with error handling."""
        try:
            paragraph.style = style
        except KeyError:
            logger.warning(
                f"Paragraph style '{style}' not found. Using default."
            )

    def _render_section(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Render a section container."""
        for child_block in data.get("children", []):
            self.render_block(child_block)

    def _render_divider(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Render a horizontal divider with proper spacing."""
        p = self._create_paragraph_at_position()

        # Set spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)

        # Add bottom border
        self._add_paragraph_border(p)

    def _add_paragraph_border(self, paragraph: Paragraph) -> None:
        """Add a bottom border to a paragraph."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
        bottom_bdr = self._create_border_element()
        p_bdr.append(bottom_bdr)

    def _render_unsupported(
        self, data: Dict[str, Any], style: Optional[str]
    ) -> None:
        """Handle unsupported block types."""
        logger.warning(f"Unsupported block type encountered. Data: {data}")

        # Add placeholder for debugging
        p = self._create_paragraph_at_position()
        run = p.add_run("[Unsupported block type - check logs for details]")
        run.italic = True
        run.font.color.rgb = Constants.COLOR_RED


class DocumentProcessor:
    """
    Main class for processing Word documents.

    This class orchestrates the entire document processing pipeline,
    from loading templates to generating final PDF output.
    """

    def __init__(
        self,
        template_path: Union[str, Path],
        output_dir: Union[str, Path] = "output",
    ):
        """
        Initialize the DocumentProcessor.

        Args:
            template_path: Path to the Word template file
            output_dir: Directory for output files (default: 'output')

        Raises:
            FileNotFoundError: If template file doesn't exist
        """
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)

        # Validate template
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

        # Create output directory
        self.output_dir.mkdir(exist_ok=True)

    def print_template_styles(self) -> None:
        """Print all available styles in the Word template for reference."""
        logger.info("\n" + "=" * 60)
        logger.info("AVAILABLE STYLES IN TEMPLATE")
        logger.info("=" * 60)

        try:
            document = Document(self.template_path)

            # Group styles by type
            style_groups = self._group_styles_by_type(document)

            # Print categorized styles
            self._print_style_groups(style_groups)

            total_styles = sum(len(styles) for styles in style_groups.values())
            logger.info(f"\n📊 TOTAL STYLES: {total_styles}")
            logger.info("=" * 60 + "\n")

        except Exception as e:
            logger.error(f"❌ Error reading template styles: {e}")
            logger.info("=" * 60 + "\n")

    @staticmethod
    def _group_styles_by_type(document: Document) -> Dict[str, List[str]]:
        """Group document styles by their type."""
        style_groups = {
            "paragraph": [],
            "character": [],
            "table": [],
            "numbering": [],
        }

        type_mapping = {
            StyleType.PARAGRAPH: "paragraph",
            StyleType.CHARACTER: "character",
            StyleType.TABLE: "table",
            StyleType.LIST: "numbering",
        }

        for style in document.styles:
            if style.type in type_mapping:
                style_groups[type_mapping[style.type]].append(style.name)

        return style_groups

    @staticmethod
    def _print_style_groups(style_groups: Dict[str, List[str]]) -> None:
        """Print grouped styles with descriptions."""
        style_info = [
            (
                "📝 PARAGRAPH STYLES",
                "paragraph",
                'Use for block-level "style" property',
            ),
            (
                "🔤 CHARACTER STYLES",
                "character",
                'Use for run-level "style" in "runs"',
            ),
            ("📊 TABLE STYLES", "table", ""),
            ("📋 LIST STYLES", "numbering", ""),
        ]

        for title, key, description in style_info:
            styles = style_groups[key]
            if styles:
                logger.info(f"\n{title} ({len(styles)}):")
                if description:
                    logger.info(f"   ({description})")
                for style in sorted(styles):
                    logger.info(f"   • {style}")

    def _extract_jinja_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract non-preformatted data for Jinja templating.

        Pre-formatted data will be handled separately by our custom renderer.

        Args:
            data: Full data dictionary

        Returns:
            Dictionary containing only non-preformatted data for Jinja
        """
        jinja_data = {}
        for key, value in data.items():
            if not self._is_dynamic_content(value):
                jinja_data[key] = value
        return jinja_data

    @staticmethod
    def _log_template_error(template_error: Exception) -> None:
        """Log detailed error message for Jinja2 template rendering failures."""
        logger.error(f"docxtpl template rendering error: {template_error}")
        logger.error("\n" + "=" * 60)
        logger.error("TEMPLATE SYNTAX ERROR")
        logger.error("=" * 60)
        logger.error(
            "The template contains syntax that is not compatible with Jinja2."
        )
        logger.error("Common issues:")
        logger.error(
            "  - Custom syntax like {{image:variable}} is not valid Jinja2"
        )
        logger.error("  - Use standard Jinja2 syntax only:")
        logger.error("    • Variables: {{variable_name}}")
        logger.error("    • Conditionals: {%if condition%}...{%endif%}")
        logger.error("    • Loops: {%for item in items%}...{%endfor%}")
        logger.error(
            "    • Table row loops: {%tr for item in items%}...{%tr endfor%}"
        )
        logger.error(
            "\nFor images, use docxtpl's InlineImage class in your data."
        )
        logger.error("=" * 60 + "\n")

    def process_document(
        self, data: Dict[str, Any], output_filename: str = "processed_document"
    ) -> Tuple[str, str]:
        """
        Main method to process the document with both pre-formatted content and Jinja templating.

        Processing order:
        1. Inject pre-formatted content blocks (dynamic structure)
        2. Process Jinja templates (variables, loops, conditionals)
        3. Convert to PDF

        Args:
            data: Dictionary containing the data to fill
            output_filename: Name for the output file (without extension)

        Returns:
            Tuple of (docx_path, pdf_path)

        Raises:
            Exception: If document processing fails
        """
        try:
            logger.info("Processing document...")

            # Step 1: Process pre-formatted content blocks FIRST
            # This injects structured content that may contain Jinja placeholders
            logger.info(
                "Step 1: Processing pre-formatted content blocks (dynamic structure)..."
            )
            document = Document(self.template_path)
            self._process_dynamic_content(document, data)

            # Save to temporary file
            temp_path = self.output_dir / f"_temp_{output_filename}.docx"
            document.save(str(temp_path))
            logger.info("Pre-formatted content injected")

            # Step 2: Use docxtpl to process all Jinja templating SECOND
            # This fills in variables, processes loops and conditionals
            logger.info(
                "Step 2: Processing Jinja templates (variables, loops, conditionals)..."
            )
            template = DocxTemplate(temp_path)

            # Separate data for Jinja (exclude pre-formatted content)
            jinja_data = self._extract_jinja_data(data)

            # Process images (convert URLs to InlineImage objects)
            logger.info("Processing images...")
            jinja_data = prepare_image_data(template, jinja_data)

            # Render with docxtpl (handles all {{variables}}, {%if%}, {%for%}, etc.)
            try:
                template.render(jinja_data)
            except Exception as template_error:
                self._log_template_error(template_error)
                raise

            # Step 3: Save the final document
            output_path = self.output_dir / f"{output_filename}.docx"
            template.save(str(output_path))
            logger.info(f"Document saved: {output_path}")

            # Clean up temp file
            if temp_path.exists():
                temp_path.unlink()

            # Step 4: Convert to PDF
            logger.info("Converting to PDF...")
            pdf_path = self._convert_to_pdf(str(output_path))

            return str(output_path), pdf_path

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise

    def _process_dynamic_content(
        self, document: Document, data: Dict[str, Any]
    ) -> None:
        """Process dynamic content (pre-formatted properties)."""
        for key, value in data.items():
            if self._is_dynamic_content(value):
                self._replace_dynamic_placeholder(document, key, value)

    @staticmethod
    def _is_dynamic_content(value: Any) -> bool:
        """Check if a value is dynamic (pre-formatted) content."""
        return (
            isinstance(value, dict)
            and value.get(Constants.DATA_TYPE_KEY)
            == Constants.PRE_FORMATTED_TYPE
        )

    def _replace_dynamic_placeholder(
        self, document: Document, key: str, value: Dict[str, Any]
    ) -> None:
        """Find and replace dynamic placeholders in the document."""
        placeholder = Constants.PLACEHOLDER_FORMAT.format(key)

        # Try to find in paragraphs first
        if self._replace_in_paragraphs(document, placeholder, value):
            return

        # Then try tables
        if self._replace_in_tables(document, placeholder, value):
            return

        logger.warning(f"Placeholder {placeholder} not found in document")

    def _replace_in_paragraphs(
        self, document: Document, placeholder: str, value: Dict[str, Any]
    ) -> bool:
        """Replace placeholder in document paragraphs."""
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                self._replace_paragraph_with_content(paragraph, value)
                return True
        return False

    def _replace_paragraph_with_content(
        self, paragraph: Paragraph, value: Dict[str, Any]
    ) -> None:
        """Replace a paragraph containing a placeholder with new content."""
        para_element = paragraph._element
        parent = para_element.getparent()
        index = parent.index(para_element)

        # Remove placeholder paragraph
        parent.remove(para_element)

        # Insert new content
        renderer = DocxRenderer(
            paragraph._parent, insert_at_parent=parent, insert_at_index=index
        )
        content = value.get("content", [])
        renderer.render_document(content)

    def _replace_in_tables(
        self, document: Document, placeholder: str, value: Dict[str, Any]
    ) -> bool:
        """Replace placeholder in document tables."""
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_cell_content(cell, value, document)
                            return True
        return False

    def _replace_cell_content(
        self, cell, value: Dict[str, Any], document: Document
    ) -> None:
        """Replace content in a table cell."""
        cell_element = cell._tc

        # Clear existing paragraphs
        for p_elem in cell_element.findall(".//w:p", cell_element.nsmap):
            cell_element.remove(p_elem)

        # Insert new content
        renderer = DocxRenderer(
            document, insert_at_parent=cell_element, insert_at_index=0
        )
        content = value.get("content", [])
        renderer.render_document(content)

    def _convert_to_pdf(self, docx_path: str) -> str:
        """
        Convert a DOCX file to PDF using UnoServer API.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Path to the generated PDF file

        Raises:
            Exception: If PDF conversion fails
        """
        import time
        from app.services.unoserver_converter import UnoServerConverter

        start_time = time.time()
        output_dir_str = str(self.output_dir)

        try:
            logger.info("Converting to PDF using UnoServer")
            converter = UnoServerConverter()

            # Check if UnoServer is available
            if not converter.is_available():
                raise Exception(
                    "UnoServer is not available. "
                    "Ensure the unoserver container is running."
                )

            # Convert using UnoServer
            pdf_path = converter.convert_to_pdf(docx_path, output_dir_str)

            conversion_time = time.time() - start_time
            logger.info(
                f"PDF conversion successful: {pdf_path} "
                f"(took {conversion_time:.2f}s)"
            )
            return pdf_path

        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")
            raise Exception(f"PDF conversion error: {e}")
